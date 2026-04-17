from typing import List, Dict, Any, Optional
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器，支持多种文件格式"""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".md"}

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """处理文件并返回分块"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        if ext == ".txt" or ext == ".md":
            return self._process_text(file_path)
        elif ext == ".pdf":
            return self._process_pdf(file_path)
        elif ext == ".docx":
            return self._process_docx(file_path)

    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """处理纯文本文件"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        chunks = self._split_text(content)

        return {
            "filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "page_count": 1,
            "content": content,
            "chunks": chunks,
        }

    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """处理 PDF 文件"""
        try:
            import PyPDF2

            chunks = []
            content_parts = []

            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                page_count = len(reader.pages)

                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
                        page_chunks = self._split_text(text)
                        for i, chunk in enumerate(page_chunks):
                            chunks.append(
                                {
                                    "content": chunk,
                                    "page": page_num + 1,
                                    "index": len(chunks),
                                }
                            )

            return {
                "filename": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "page_count": page_count,
                "content": "\n\n".join(content_parts),
                "chunks": chunks,
            }
        except ImportError:
            logger.warning("PyPDF2 not installed, using fallback text extraction")
            return self._process_text_fallback(file_path)

    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """处理 DOCX 文件"""
        try:
            from docx import Document

            content_parts = []
            chunks = []

            doc = Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            content = "\n\n".join(full_text)
            content_parts.append(content)
            text_chunks = self._split_text(content)

            for i, chunk in enumerate(text_chunks):
                chunks.append(
                    {
                        "content": chunk,
                        "page": 1,
                        "index": i,
                    }
                )

            return {
                "filename": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "page_count": len(doc.paragraphs) or 1,
                "content": content,
                "chunks": chunks,
            }
        except ImportError:
            logger.warning("python-docx not installed")
            raise ImportError("Please install python-docx: pip install python-docx")

    def _process_text_fallback(self, file_path: str) -> Dict[str, Any]:
        """PDF 文本提取失败的回退处理"""
        return {
            "filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "page_count": 1,
            "content": "[PDF content could not be extracted]",
            "chunks": [
                {
                    "content": "[PDF content could not be extracted]",
                    "page": 1,
                    "index": 0,
                }
            ],
        }

    def _split_text(self, text: str) -> List[str]:
        """将文本分割成块"""
        if not text or not text.strip():
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + self.chunk_size

            if end < text_length:
                separator = self._find_separator(text, start, end)
                if separator != -1:
                    end = separator + 1

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = (
                end - self.chunk_overlap if end - self.chunk_overlap > start else end
            )

        return chunks

    def _find_separator(self, text: str, start: int, end: int) -> int:
        """在指定范围内查找合适的分割点"""
        separators = ["\n\n", "\n", ". ", "。", "! ", "！", "? ", "？"]

        best_pos = -1
        for sep in separators:
            pos = text.rfind(sep, start, end)
            if pos != -1:
                if best_pos == -1 or pos > best_pos:
                    best_pos = pos

        return best_pos


def get_document_processor(
    chunk_size: int = 500, chunk_overlap: int = 50
) -> DocumentProcessor:
    """获取文档处理器实例"""
    return DocumentProcessor(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
